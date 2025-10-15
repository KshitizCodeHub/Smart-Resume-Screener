"""
PDF and DOCX parsing service.
Extracts text content from resume files.
"""
import pdfplumber
from docx import Document
from typing import Optional
import io

class DocumentParser:
    """Parse PDF and DOCX files to extract text content."""
    
    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file.
        Uses pdfplumber for better text extraction.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text as string
        """
        try:
            text = ""
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error parsing PDF with pdfplumber: {e}")
            # Fallback to pypdf (lazy import to avoid xml.dom.NodeFilter bug)
            try:
                from pypdf import PdfReader
                pdf_reader = PdfReader(io.BytesIO(file_content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
            except Exception as e2:
                print(f"pypdf fallback also failed: {e2}")
                raise Exception("Failed to parse PDF file")
    
    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text as string
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return text.strip()
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            raise Exception("Failed to parse DOCX file")
    
    @staticmethod
    def parse_txt(file_content: bytes) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_content: TXT file content as bytes
            
        Returns:
            Extracted text as string
        """
        try:
            return file_content.decode('utf-8').strip()
        except UnicodeDecodeError:
            # Try other encodings
            try:
                return file_content.decode('latin-1').strip()
            except Exception as e:
                print(f"Error parsing TXT: {e}")
                raise Exception("Failed to parse TXT file")
    
    @staticmethod
    def parse_file(filename: str, file_content: bytes) -> str:
        """
        Parse file based on extension.
        
        Args:
            filename: Name of the file
            file_content: File content as bytes
            
        Returns:
            Extracted text as string
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return DocumentParser.parse_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            return DocumentParser.parse_docx(file_content)
        elif filename_lower.endswith('.txt'):
            return DocumentParser.parse_txt(file_content)
        else:
            raise Exception(f"Unsupported file format: {filename}")
